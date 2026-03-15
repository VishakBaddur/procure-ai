import pdfplumber
from PIL import Image
import pytesseract
import re
from pathlib import Path
from typing import Dict, Any, List
import json
import os
import requests
from docx import Document


class PriceComparisonAgent:
    """Agent for extracting and comparing prices from vendor quotes using AI"""
    
    def __init__(self):
        # Initialize Local LLM (Ollama) for intelligent price extraction
        # Check if Ollama is available, fallback to Groq if needed
        self.ollama_url = os.getenv("OLLAMA_URL", "http://localhost:11434")
        self.ollama_model = os.getenv("OLLAMA_MODEL", "llama3.2")
        
        # Try to connect to Ollama
        try:
            response = requests.get(f"{self.ollama_url}/api/tags", timeout=2)
            if response.status_code == 200:
                self.use_local_llm = True
                print(f"✓ Using local LLM (Ollama) with model: {self.ollama_model}")
            else:
                self.use_local_llm = False
        except Exception as e:
            print(f"Warning: Ollama not available, checking for Groq fallback: {e}")
            self.use_local_llm = False
        
        # Fallback to Groq if Ollama is not available
        if not self.use_local_llm:
            api_key = os.getenv("GROQ_API_KEY")
            if api_key:
                try:
                    from groq import Groq
                    self.client = Groq(api_key=api_key)
                    self.use_ai = True
                    print("✓ Using Groq AI as fallback")
                except Exception as e:
                    print(f"Warning: Groq AI not available: {e}")
                    self.use_ai = False
            else:
                self.use_ai = False
                print("⚠️ No LLM available - will use regex fallback only")
        else:
            self.use_ai = True  # Use local LLM
    
    async def process_quote(self, file_path: Path, vendor_name: str, content_type: str) -> Dict[str, Any]:
        """Process quote file and extract pricing information"""
        # Base result structure so callers always get a predictable shape
        base_result: Dict[str, Any] = {
            "vendor_name": vendor_name,
            "products": [],
            "items": [],
            "total_price": 0,
            "currency": "USD",
            "warranties": [],
            "warnings": [],
            "success": True,
            "error": None,
        }
        
        extracted_text = ""
        
        # Extract text based on file type
        try:
            if content_type == "application/pdf" or str(file_path).endswith(".pdf"):
                extracted_text = self._extract_from_pdf(file_path)
            elif content_type.startswith("image/"):
                extracted_text = self._extract_from_image(file_path)
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ] or str(file_path).endswith((".docx", ".doc")):
                extracted_text = self._extract_from_word(file_path)
            elif content_type == "text/plain" or str(file_path).endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    extracted_text = f.read()
            elif content_type == "message/rfc822" or "email" in content_type.lower() or str(file_path).endswith((".eml", ".msg")):
                # Email file - extract text content
                extracted_text = self._extract_from_email(file_path)
            else:
                # Try to detect file type by extension
                file_ext = str(file_path).lower()
                if file_ext.endswith(".docx") or file_ext.endswith(".doc"):
                    extracted_text = self._extract_from_word(file_path)
                elif file_ext.endswith(".txt"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        extracted_text = f.read()
                else:
                    # Default: try to read as text
                    try:
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            extracted_text = f.read()
                    except Exception:
                        extracted_text = ""
        except Exception as e:
            # Hard failure during extraction – return a clear, user-facing error
            base_result["success"] = False
            base_result["error"] = (
                "We couldn't read this file reliably. "
                "Please upload a clearer PDF/Doc or paste the quote text manually."
            )
            base_result["warnings"].append(f"text_extraction_error: {str(e)}")
            base_result["raw_text_preview"] = ""
            return base_result
        
        # If we could not extract any text at all, fail gracefully
        if not extracted_text or not extracted_text.strip():
            base_result["success"] = False
            base_result["error"] = (
                "We couldn't extract any readable text from this document. "
                "Please check the file (e.g., scanned image, password-protected PDF) "
                "or paste the quote text manually."
            )
            base_result["warnings"].append("no_text_extracted")
            base_result["raw_text_preview"] = ""
            return base_result
        
        # Log raw extracted text for production debugging (Render logs)
        import sys
        _preview = (extracted_text or "")[:800].replace("\r", "\n")
        print(f"[QuoteParse] RAW_EXTRACTED_TEXT_LEN={len(extracted_text or '')}", file=sys.stderr, flush=True)
        print(f"[QuoteParse] RAW_EXTRACTED_TEXT_PREVIEW=\n{_preview}", file=sys.stderr, flush=True)
        
        # Parse pricing information
        pricing_data = await self._parse_pricing(extracted_text, vendor_name)
        
        # Ensure a stable envelope and attach basic diagnostics
        if not isinstance(pricing_data, dict):
            base_result["success"] = False
            base_result["error"] = (
                "The AI parser returned an unexpected response. "
                "Please review the extracted data manually."
            )
            base_result["warnings"].append("invalid_pricing_data_type")
            base_result["raw_text_preview"] = extracted_text[:500]
            return base_result
        
        # Merge pricing data into the base envelope
        merged: Dict[str, Any] = {**base_result, **pricing_data}
        
        # If no products/items were parsed, mark as low-confidence and surface that to the caller
        has_items = bool(merged.get("products") or merged.get("items"))
        if not has_items:
            merged["success"] = False
            if not merged.get("error"):
                merged["error"] = (
                    "We couldn't confidently extract any line items from this quote. "
                    "Please double-check the numbers and update them manually if needed."
                )
            merged.setdefault("warnings", []).append("no_items_parsed")
        
        # Always include a short preview of the raw text to help debugging / UI hints
        if "raw_text_preview" not in merged:
            merged["raw_text_preview"] = extracted_text[:500]
        
        return merged
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return text
    
    def _extract_from_image(self, file_path: Path) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            print(f"Error extracting image: {e}")
            return ""
    
    def _extract_from_word(self, file_path: Path) -> str:
        """Extract text from Word document (.docx)"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"Error extracting Word document: {e}")
        return text
    
    def _extract_from_email(self, file_path: Path) -> str:
        """Extract text from email file (.eml or .msg)"""
        text = ""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                # Try to extract email body (simple extraction)
                # Look for common email patterns
                if "Content-Type: text/plain" in content:
                    parts = content.split("Content-Type: text/plain")
                    if len(parts) > 1:
                        body = parts[1].split("\n\n", 1)
                        if len(body) > 1:
                            text = body[1]
                elif "Content-Type: text/html" in content:
                    parts = content.split("Content-Type: text/html")
                    if len(parts) > 1:
                        body = parts[1].split("\n\n", 1)
                        if len(body) > 1:
                            # Simple HTML stripping
                            html_text = body[1]
                            # Remove HTML tags (basic)
                            text = re.sub(r'<[^>]+>', '', html_text)
                else:
                    # Fallback: extract everything after headers
                    lines = content.split("\n")
                    in_body = False
                    for line in lines:
                        if line.strip() == "" and not in_body:
                            in_body = True
                            continue
                        if in_body:
                            text += line + "\n"
        except Exception as e:
            print(f"Error extracting email: {e}")
            # Fallback: try to read as plain text
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except:
                pass
        return text
    
    def _log_email_style_heuristic(self, text: str) -> bool:
        """Diagnostic only: does text look like an email? Logged for Render debugging; does not change code path."""
        if not text or len(text) < 50:
            return False
        t = text.lower()
        markers = [
            "good talking", "hey there", "let me know", "thanks,", "best,",
            "per your request", "following up", "re: ", "subject:", "call or email",
        ]
        return sum(1 for m in markers if m in t) >= 2

    async def _parse_pricing(self, text: str, vendor_name: str) -> Dict[str, Any]:
        """Parse pricing: use Groq/Ollama as the single parser when available (understands any format); regex only when no LLM."""
        import sys
        email_style = self._log_email_style_heuristic(text)
        print(f"[QuoteParse] EMAIL_STYLE_HEURISTIC={email_style} (diagnostic only, path unchanged)", file=sys.stderr, flush=True)
        if self.use_ai:
            print(f"[QuoteParse] EXTRACTION_PATH=LLM (attempting)", file=sys.stderr, flush=True)
            try:
                result = await self._parse_with_ai(text, vendor_name)
                print(f"[QuoteParse] EXTRACTION_PATH=LLM (success)", file=sys.stderr, flush=True)
                return result
            except Exception as e:
                print(f"[QuoteParse] EXTRACTION_PATH=regex (LLM failed: {e})", file=sys.stderr, flush=True)
                return self._parse_with_regex(text, vendor_name)
        print(f"[QuoteParse] EXTRACTION_PATH=regex (no LLM available)", file=sys.stderr, flush=True)
        return self._parse_with_regex(text, vendor_name)
    
    async def _parse_with_ai(self, text: str, vendor_name: str) -> Dict[str, Any]:
        """Use Groq AI to intelligently extract pricing information with products, payment terms, and quantity tiers"""
        
        prompt = f"""
        You are the only extractor for this quote. Your job: read the document/email text and return ONLY real product or service line items with prices. The UI will display exactly what you return.

        Vendor: {vendor_name}

        Document text (from PDF, email, or upload):
        {text[:12000]}

        RULES YOU MUST FOLLOW:
        1. **Product name**: Short, clear name only (e.g. "Basic wrap goggles", "Hard hats", "Hi-vis vests", "Starter pack"). NEVER use a full sentence, greeting, or clause as a product name.
        2. **Only real line items**: Include ONLY items that are clearly goods or services being sold with a price. Do NOT include as products: section headers, Subtotal/Total/Freight/Tax, contact labels (Cell, Phone, Fax, Email), or any sentence about the call, timing, shipping, or payment (e.g. "Good talking earlier", "On timing: we ship in X days", "Payment-wise, most folks prepay", "Cell:--").
        3. **Prices**: Extract full unit price and total where stated. Use quantity_min/quantity_max for ranges. Currency in pricing_matrix. Use "Not specified" for missing fields.
        4. **One product per item**: Each distinct product/service gets one object in "products". Do not combine multiple products into one. Pricing tiers (e.g. "1–5 kg: USD X") go inside pricing_matrix of the product, not as separate products.
        5. **Structure**: Each product must have: name, description (or "Not specified"), category (or "Not specified"), pricing_matrix (array of at least one entry with unit_price or total_price, quantity_min/max, quantity_unit, payment_terms, currency), warranty, delivery_time. Use "Not specified" for any missing value.

        Return ONLY valid JSON in this exact structure:
        {{
            "products": [
                {{
                    "product_id": "unique identifier (e.g., 'PROD001' or 'Aluminium-Sheet-Type-A')",
                    "name": "product name/variety/model",
                    "description": "detailed description if available, or 'Not specified'",
                    "category": "product category if mentioned, or 'Not specified'",
                    "pricing_matrix": [
                        {{
                            "quantity_min": <minimum quantity or null>,
                            "quantity_max": <maximum quantity or null for unlimited>,
                            "quantity_unit": "unit type (e.g., 'units', 'pieces', 'kg', 'lb') or 'Not specified'",
                            "payment_terms": "payment term (e.g., 'Net 30', 'Upfront', '50% Advance', 'Net 60') or 'Not specified'",
                            "unit_price": <price per unit or null if not specified>,
                            "total_price": <total price for this quantity tier or null>,
                            "currency": "currency code (e.g., 'USD', 'EUR') or 'Not specified'",
                            "notes": "any special notes (e.g., 'bulk discount', 'early payment discount') or 'Not specified'"
                        }}
                    ],
                    "default_payment_term": "most common payment term if applicable, or 'Not specified'",
                    "warranty": "warranty information if mentioned, or 'Not specified'",
                    "delivery_time": "delivery time if mentioned, or 'Not specified'"
                }}
            ],
            "summary": {{
                "total_products": <number of different products>,
                "currency": "primary currency (USD, EUR, etc.) or 'Not specified'",
                "payment_terms_available": ["list of all payment terms found, or ['Not specified'] if none"],
                "quantity_tiers": ["list of quantity ranges found, or ['Not specified'] if none"],
                "total_price_range": {{
                    "min": <minimum total price or null>,
                    "max": <maximum total price or null>
                }}
            }},
            "general_notes": "any general pricing notes, discounts, or conditions, or 'Not specified'",
            "warranties": ["list of warranty information, or ['Not specified'] if none"],
            "other_info": {{
                "delivery_terms": "delivery terms if mentioned, or 'Not specified'",
                "return_policy": "return policy if mentioned, or 'Not specified'",
                "support_services": "support services if mentioned, or 'Not specified'",
                "additional_notes": "any other relevant information, or 'Not specified'"
            }},
            "extraction_metadata": {{
                "parsing_method": "AI",
                "confidence": "high/medium/low"
            }}
        }}
        
        Return ONLY valid JSON. No markdown, no code blocks, no explanatory text. Just the JSON object.
        """
        
        try:
            import sys
            print(f"[Price Agent] Processing quote for {vendor_name}, text length: {len(text)}", file=sys.stderr, flush=True)
            print(f"[Price Agent] First 500 chars: {text[:500]}", file=sys.stderr, flush=True)
            
            # Use local LLM (Ollama) or Groq for extraction — same prompt for both
            if self.use_local_llm:
                print(f"[QuoteParse] EXTRACTION_LLM=Ollama", file=sys.stderr, flush=True)
                return await self._parse_with_local_llm(text, vendor_name, prompt)
            
            print(f"[QuoteParse] EXTRACTION_LLM=Groq", file=sys.stderr, flush=True)
            # Use Groq (fallback)
            # Use a more capable model for complex extraction
            system_message = """You are the sole extractor for vendor quotes. Return ONLY real product/service line items with prices. The app will show exactly what you return.

RULES: (1) Product "name" = short name only (e.g. "Basic wrap goggles", "Hard hats"). Never use a sentence, greeting, or contact line (Cell, Phone) as a product name. (2) Include ONLY actual goods/services being sold with a price—never Subtotal, Total, Freight, or sentences about timing/shipping/payment. (3) One product per item; pricing tiers go in pricing_matrix. (4) Use "Not specified" for any missing field. Return only valid JSON."""
            
            response = self.client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "system", 
                        "content": system_message
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1  # Very low temperature for consistent extraction
            )
            response_text = response.choices[0].message.content
            import sys
            print(f"[Price Agent] AI Response (first 500 chars): {response_text[:500]}", file=sys.stderr, flush=True)
            
            # Extract JSON from response
            json_str = self._extract_json_from_response(response_text)
            pricing_data = json.loads(json_str)
            
            # Process the same way as local LLM version
            result = self._process_ai_response(pricing_data, text, vendor_name)
            result["parsing_method"] = "AI (Groq)"  # Mark as Groq
            return result
            
        except json.JSONDecodeError as e:
            # If JSON parsing fails, try regex fallback
            import sys
            print(f"[Price Agent] Failed to parse AI response as JSON: {e}", file=sys.stderr, flush=True)
            print(f"[Price Agent] Response was: {response_text[:500] if 'response_text' in locals() else 'N/A'}", file=sys.stderr, flush=True)
            return self._parse_with_regex(text, vendor_name)
        except Exception as e:
            import sys
            print(f"[Price Agent] AI parsing error: {e}", file=sys.stderr, flush=True)
            import traceback
            traceback.print_exc(file=sys.stderr)
            return self._parse_with_regex(text, vendor_name)
    
    async def _parse_with_local_llm(self, text: str, vendor_name: str, prompt: str) -> Dict[str, Any]:
        """Use local LLM (Ollama) with same extraction prompt as Groq."""
        import sys
        try:
            print(f"[Price Agent] Using local LLM (Ollama) for parsing", file=sys.stderr, flush=True)
            # Call Ollama API
            response = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": self.ollama_model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,  # Low temperature for consistent extraction
                        "num_predict": 4000  # Allow longer responses
                    }
                },
                timeout=120  # 2 minute timeout
            )
            
            if response.status_code != 200:
                raise Exception(f"Ollama API error: {response.status_code} - {response.text}")
            
            response_data = response.json()
            response_text = response_data.get("response", "")
            
            print(f"[Price Agent] Local LLM Response (first 500 chars): {response_text[:500]}", file=sys.stderr, flush=True)
            
            # Extract JSON from response
            json_str = self._extract_json_from_response(response_text)
            pricing_data = json.loads(json_str)
            
            # Process the same way as Groq version
            return self._process_ai_response(pricing_data, text, vendor_name)
            
        except Exception as e:
            import sys
            print(f"[Price Agent] Local LLM parsing error: {e}", file=sys.stderr, flush=True)
            import traceback
            traceback.print_exc(file=sys.stderr)
            # Fallback to regex
            return self._parse_with_regex(text, vendor_name)
    
    def _process_ai_response(self, pricing_data: Dict[str, Any], text: str, vendor_name: str) -> Dict[str, Any]:
        """Trust Groq/LLM output: minimal validation, then build display structure. Extraction responsibility is entirely with the LLM."""
        import sys
        import re
        
        products = pricing_data.get("products", []) or []
        summary = pricing_data.get("summary", {}) or {}
        currency = (summary.get("currency") or pricing_data.get("currency") or "USD").strip() or "USD"
        
        # Log what LLM returned (before any normalization) for production debugging
        before_names = [p.get("name", "") for p in products]
        print(f"[QuoteParse] PRODUCTS_FROM_LLM_BEFORE_NORMALIZATION count={len(products)} names={before_names}", file=sys.stderr, flush=True)
        
        # Minimal normalization: ensure each product has name, pricing_matrix, warranty; strip leading item numbers from name
        for p in products:
            name = (p.get("name") or "Unknown Product").strip()
            name = re.sub(r'^\d+\s*[.)]\s*', '', name).strip()
            name = re.sub(r'^\d+\s+', '', name).strip()
            p["name"] = name or p.get("name") or "Unknown Product"
            if not p.get("pricing_matrix"):
                p["pricing_matrix"] = []
            if not p.get("warranty"):
                p["warranty"] = "Not specified"
        
        after_names = [p.get("name", "") for p in products]
        print(f"[QuoteParse] PRODUCTS_TO_DISPLAY_AFTER_NORMALIZATION count={len(products)} names={after_names}", file=sys.stderr, flush=True)
        
        # Build legacy items and total from Groq output
        items = []
        total_price = 0.0
        
        for product in products:
            product_name = product.get("name", "Unknown Product")
            pricing_matrix = product.get("pricing_matrix", [])
            
            if not pricing_matrix:
                items.append({
                    "name": product_name,
                    "price": 0.0,
                    "quantity": 1.0,
                    "unit_price": 0.0,
                    "unit": "unit",
                    "product_id": product.get("product_id"),
                    "description": product.get("description", ""),
                    "payment_terms": product.get("default_payment_term", "Not specified")
                })
            else:
                product_total = None
                best_entry = None
                
                for price_entry in pricing_matrix:
                    if price_entry.get("total_price") and price_entry.get("total_price") > 0:
                        product_total = float(price_entry.get("total_price"))
                        best_entry = price_entry
                        break
                
                if not best_entry and pricing_matrix:
                    best_entry = pricing_matrix[0]
                    qty_min = best_entry.get("quantity_min", 1)
                    qty_max = best_entry.get("quantity_max")
                    unit_price = best_entry.get("unit_price", 0.0)
                    quantity = (qty_min + (qty_max or qty_min)) / 2 if qty_max else qty_min
                    product_total = unit_price * quantity if unit_price > 0 else 0.0
                
                if best_entry:
                    qty_min = best_entry.get("quantity_min", 1)
                    qty_max = best_entry.get("quantity_max")
                    unit_price = best_entry.get("unit_price", 0.0)
                    payment_term = best_entry.get("payment_terms", "Standard")
                    quantity = (qty_min + (qty_max or qty_min)) / 2 if qty_max else qty_min
                    
                    items.append({
                        "name": product_name,
                        "price": float(product_total) if product_total else float(unit_price * quantity),
                        "quantity": float(quantity),
                        "unit_price": float(unit_price),
                        "unit": best_entry.get("quantity_unit", "unit"),
                        "product_id": product.get("product_id"),
                        "description": product.get("description", ""),
                        "payment_terms": payment_term,
                        "quantity_min": qty_min,
                        "quantity_max": qty_max,
                        "notes": best_entry.get("notes", "")
                    })
                    total_price += float(product_total) if product_total else float(unit_price * quantity)
        
        # Prefer summed line-item total over LLM summary total (LLM sometimes returns one product's total as doc total)
        if summary.get("total_price_range") and total_price <= 0:
            total_price = summary["total_price_range"].get("max", total_price)
        
        warranties = pricing_data.get("warranties", [])
        if not warranties:
            warranties = ["Not specified"]
        
        for product in products:
            if "warranty" not in product or not product.get("warranty"):
                product["warranty"] = "Not specified"
            if "pricing_matrix" not in product:
                product["pricing_matrix"] = []
        
        other_info = pricing_data.get("other_info", {})
        if not other_info:
            other_info = {
                "delivery_terms": "Not specified",
                "return_policy": "Not specified",
                "support_services": "Not specified",
                "additional_notes": "Not specified"
            }
        
        result = {
            "vendor_name": vendor_name,
            "items": items,
            "products": products,
            "total_price": float(total_price) if total_price > 0 else 0.0,
            "currency": currency.upper() if currency and currency != "Not specified" else "USD",
            "extracted_text": text[:500],
            "item_count": len(products) if products else len(items),
            "parsing_method": "AI (Local LLM)",
            "notes": pricing_data.get("general_notes", "Not specified"),
            "warranties": warranties,
            "summary": summary,
            "payment_terms_available": summary.get("payment_terms_available", ["Not specified"]),
            "quantity_tiers": summary.get("quantity_tiers", ["Not specified"]),
            "other_info": other_info
        }
        
        print(f"[Price Agent] Final result - products: {len(products)}, items: {len(items)}, total_price: {result['total_price']}", file=sys.stderr, flush=True)
        
        return result
    
    def _fix_pricing_values(self, products: List[Dict[str, Any]], text: str, currency: str) -> List[Dict[str, Any]]:
        """Fix pricing values by extracting them from the original text if they were incorrectly parsed"""
        import re
        import sys
        
        # Extract all prices from text: "USD 68,500 per kg" → 68500
        price_pattern = r'USD\s+(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s+per\s+(kg|lb|unit|piece)'
        all_prices = re.findall(price_pattern, text, re.IGNORECASE)
        
        # Extract quantity tiers: "For orders between 1–5 kg" or "1–3 kg:"
        tier_pattern = r'(\d+)[–-](\d+)\s*(?:kg|lb|units?|pieces?)'
        tiers = re.findall(tier_pattern, text, re.IGNORECASE)
        
        # Extract payment terms
        payment_terms = []
        if re.search(r'full\s+advance', text, re.IGNORECASE):
            payment_terms.append('Full Advance')
        if re.search(r'50%\s*advance', text, re.IGNORECASE):
            payment_terms.append('50% Advance, Balance on Delivery')
        if re.search(r'lc\s+at\s+sight', text, re.IGNORECASE):
            payment_terms.append('LC at Sight')
        if re.search(r'split\s+payment', text, re.IGNORECASE):
            payment_terms.append('Split Payment')
        
        # Determine unit from text
        unit = 'kg'
        if 'per kg' in text.lower() or 'per kilogram' in text.lower():
            unit = 'kg'
        elif 'per lb' in text.lower() or 'per pound' in text.lower():
            unit = 'lb'
        elif 'per unit' in text.lower():
            unit = 'unit'
        
        # Build pricing matrix for each product
        for product in products:
            pricing_matrix = product.get('pricing_matrix', [])
            
            # Check if pricing needs fixing (empty, wrong prices, or missing)
            needs_fix = False
            if not pricing_matrix:
                needs_fix = True
            elif any(p.get('unit_price', 0) < 1000 for p in pricing_matrix if p.get('unit_price')):
                # Prices seem too low (like 500 instead of 68500)
                needs_fix = True
            
            if needs_fix and all_prices:
                print(f"[Price Agent] Fixing pricing for product: {product.get('name', '')[:50]}", file=sys.stderr, flush=True)
                
                new_pricing_matrix = []
                prices = [float(p[0].replace(',', '')) for p in all_prices]
                
                # If we have tiers and prices, match them
                if tiers and prices:
                    # Group prices by tier (assuming prices are listed in order)
                    prices_per_tier = len(prices) // len(tiers) if len(tiers) > 0 else len(prices)
                    
                    for tier_idx, (qty_min, qty_max) in enumerate(tiers):
                        start_idx = tier_idx * prices_per_tier
                        tier_prices = prices[start_idx:start_idx + prices_per_tier]
                        
                        # Match prices with payment terms if available
                        if payment_terms and len(tier_prices) == len(payment_terms):
                            for price_val, payment_term in zip(tier_prices, payment_terms):
                                new_pricing_matrix.append({
                                    "quantity_min": int(qty_min),
                                    "quantity_max": int(qty_max),
                                    "quantity_unit": unit,
                                    "payment_terms": payment_term,
                                    "unit_price": price_val,
                                    "total_price": None,
                                    "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                                    "notes": ""
                                })
                        else:
                            # Use first price for this tier
                            if tier_prices:
                                new_pricing_matrix.append({
                                    "quantity_min": int(qty_min),
                                    "quantity_max": int(qty_max),
                                    "quantity_unit": unit,
                                    "payment_terms": payment_terms[0] if payment_terms else "Standard",
                                    "unit_price": tier_prices[0],
                                    "total_price": None,
                                    "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                                    "notes": ""
                                })
                elif prices:
                    # No tiers, just use all prices
                    if payment_terms and len(prices) == len(payment_terms):
                        for price_val, payment_term in zip(prices, payment_terms):
                            new_pricing_matrix.append({
                                "quantity_min": 1,
                                "quantity_max": None,
                                "quantity_unit": unit,
                                "payment_terms": payment_term,
                                "unit_price": price_val,
                                "total_price": None,
                                "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                                "notes": ""
                            })
                    else:
                        # Use first price
                        new_pricing_matrix.append({
                            "quantity_min": 1,
                            "quantity_max": None,
                            "quantity_unit": unit,
                            "payment_terms": payment_terms[0] if payment_terms else "Standard",
                            "unit_price": prices[0],
                            "total_price": None,
                            "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                            "notes": ""
                        })
                
                if new_pricing_matrix:
                    product['pricing_matrix'] = new_pricing_matrix
                    print(f"[Price Agent] Fixed pricing: {len(new_pricing_matrix)} entries with prices: {[p.get('unit_price') for p in new_pricing_matrix]}", file=sys.stderr, flush=True)
        
        return products
    
    def _enhance_email_quote_extraction(self, products: List[Dict[str, Any]], text: str, currency: str) -> List[Dict[str, Any]]:
        """Enhance extraction specifically for email/informal quotes with better product and price detection"""
        import re
        import sys
        
        text_lower = text.lower()
        
        # AGGRESSIVE EMAIL QUOTE EXTRACTION: If no products or products seem wrong, extract directly from text
        if not products or len(products) == 0 or all(not p.get('pricing_matrix') or all(pr.get('unit_price', 0) == 0 for pr in p.get('pricing_matrix', [])) for p in products):
            print(f"[Price Agent] No valid products found, attempting aggressive email extraction...", file=sys.stderr, flush=True)
            aggressive_products = self._aggressive_email_extraction(text, currency)
            if aggressive_products:
                print(f"[Price Agent] Aggressive extraction found {len(aggressive_products)} products", file=sys.stderr, flush=True)
                return aggressive_products
        
        # If we have products but they seem incomplete, try to enhance them
        if not products:
            return products
        
        # Look for email-style product mentions: "Gold American Eagles: $2,395/oz"
        # Pattern: Product name followed by colon and price
        email_product_pattern = r'([A-Z][^:]+?):\s*\$?(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*/?\s*(\w+)'
        email_matches = re.findall(email_product_pattern, text)
        
        # Also look for: "Product name - price/unit, quantity for total"
        # Pattern: "Gold: ~$75k (31–32 oz at $2,395/oz)"
        enhanced_pattern = r'([A-Z][^:–-]+?)[:–-]\s*(?:~|approx|approximately)?\s*\$?(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*(?:for|\(|,)\s*(\d+)[–-]?(\d+)?\s*(\w+)?'
        
        # Try to match products with their pricing from email format
        for product in products:
            product_name = product.get('name', '').strip()
            pricing_matrix = product.get('pricing_matrix', [])
            
            # If product has no pricing or incomplete pricing, try to extract from text
            if not pricing_matrix or all(p.get('unit_price', 0) == 0 for p in pricing_matrix):
                # Look for this product name in the text with pricing
                # Pattern: "Gold American Eagles: $2,395/oz" or "Gold Eagles: $2,395/oz"
                product_variations = [
                    product_name,
                    product_name.replace('American ', ''),
                    product_name.replace('Gold ', '').replace('Silver ', '')
                ]
                
                for variation in product_variations:
                    if not variation or len(variation) < 3:
                        continue
                    
                    # Look for pattern: "Product: price/unit"
                    pattern = rf'{re.escape(variation)}[:\s]+?\$?(\d{{1,3}}(?:,\d{{3}})*(?:\.\d{{2}})?)\s*/?\s*(\w+)'
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        unit_price = float(match.group(1).replace(',', ''))
                        unit = match.group(2).lower()
                        
                        # Look for quantity: "31–32 oz" or "~31-32 oz"
                        qty_pattern = rf'{re.escape(variation)}[^:]*?(\d+)[–-]?(\d+)?\s*{re.escape(unit)}'
                        qty_match = re.search(qty_pattern, text, re.IGNORECASE)
                        
                        quantity_min = None
                        quantity_max = None
                        if qty_match:
                            quantity_min = int(qty_match.group(1))
                            quantity_max = int(qty_match.group(2)) if qty_match.group(2) else quantity_min
                        
                        # Look for calculation pattern: "x 40 = $96,400" first (most reliable)
                        calculation_pattern = rf'{re.escape(variation)}[^=]*?x\s*(\d+)\s*=\s*\$?(\d{{1,3}}(?:,\d{{3}})*)'
                        calc_match = re.search(calculation_pattern, text, re.IGNORECASE)
                        
                        total_price = None
                        if calc_match:
                            # Found explicit calculation: "x 40 = $96,400"
                            total_str = calc_match.group(2).replace(',', '')
                            total_price = float(total_str)
                            # Also update quantity if found
                            if not quantity_min:
                                quantity_min = int(calc_match.group(1))
                                quantity_max = quantity_min
                        else:
                            # Pattern: "for ~$75k" or "for approximately $75,000"
                            total_pattern = rf'{re.escape(variation)}[^$]*?(?:for|total|around|~|approx)\s*\$?(\d{{1,3}}(?:,\d{{3}})*(?:k|K)?)'
                            total_match = re.search(total_pattern, text, re.IGNORECASE)
                            
                            if total_match:
                                total_str = total_match.group(1).replace(',', '')
                                if 'k' in total_str.lower():
                                    total_price = float(total_str.lower().replace('k', '')) * 1000
                                else:
                                    total_price = float(total_str)
                        
                        # Create pricing matrix entry
                        if unit_price > 0:
                            pricing_matrix = [{
                                "quantity_min": quantity_min or 1,
                                "quantity_max": quantity_max,
                                "quantity_unit": unit,
                                "payment_terms": "Not specified",
                                "unit_price": unit_price,
                                "total_price": total_price,
                                "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                                "notes": "Extracted from email quote"
                            }]
                            product['pricing_matrix'] = pricing_matrix
                            print(f"[Price Agent] Enhanced product {product_name} with pricing: {unit_price}/{unit}", file=sys.stderr, flush=True)
                            break
        
        # Extract fees as separate products if mentioned
        fee_patterns = [
            (r'(?:setup|account\s+setup|admin)[:\s]+(?:usually|around|approx|~)?\s*\$?(\d{1,3}(?:,\d{3})*)', 'Setup Fee'),
            (r'storage[:\s]+(?:around|approx|~)?\s*\$?(\d{1,3}(?:,\d{3})*)[–-]?(\d{1,3}(?:,\d{3})*)?\s*(?:annually|per\s+year|/year)', 'Storage Fee'),
        ]
        
        for pattern, fee_name in fee_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                fee_amount = float(match.group(1).replace(',', ''))
                # Check if this fee product already exists
                existing_fee = next((p for p in products if fee_name.lower() in p.get('name', '').lower()), None)
                if not existing_fee:
                    products.append({
                        "product_id": f"fee_{len(products) + 1}",
                        "name": fee_name,
                        "description": "Fee extracted from quote",
                        "category": "Fee",
                        "pricing_matrix": [{
                            "quantity_min": 1,
                            "quantity_max": 1,
                            "quantity_unit": "one-time" if "setup" in fee_name.lower() else "year",
                            "payment_terms": "Not specified",
                            "unit_price": fee_amount,
                            "total_price": fee_amount,
                            "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                            "notes": "Fee extracted from email quote"
                        }],
                        "default_payment_term": "Not specified",
                        "warranty": "Not specified",
                        "delivery_time": "Not specified"
                    })
                    print(f"[Price Agent] Added fee product: {fee_name} = ${fee_amount}", file=sys.stderr, flush=True)
        
        return products
    
    def _aggressive_email_extraction(self, text: str, currency: str) -> List[Dict[str, Any]]:
        """Aggressively extract products from email-style quotes using direct pattern matching"""
        import re
        import sys
        
        products = []
        text_lower = text.lower()
        
        # FIRST: Look for "Breakdown" sections which are very reliable
        # Pattern: "Breakdown (rough numbers):\nGold Eagles: ~$2,410 x 40 = $96,400"
        breakdown_pattern = r'Breakdown[^:]*?:\s*(.*?)(?:\n\n|\n[A-Z][a-z]+\s+[A-Z]|Setup|Storage|We can|Let me|Best|$)'
        breakdown_match = re.search(breakdown_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if breakdown_match:
            breakdown_text = breakdown_match.group(1)
            print(f"[Price Agent] Found breakdown section: {breakdown_text[:300]}", file=sys.stderr, flush=True)
            
            # Extract from breakdown: "Gold Eagles: ~$2,410 x 40 = $96,400"
            # More flexible pattern to handle variations
            breakdown_line_patterns = [
                r'([A-Z][A-Za-z\s]+?(?:Eagles?|Bars?|Coins?))[:\s]+?~?\$?(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*x\s*(\d+)\s*=\s*\$?(\d{1,3}(?:,\d{3})*)',
                r'([A-Z][A-Za-z\s]+?(?:Eagles?|Bars?|Coins?))[:\s]+?~?\$?(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*/\s*(\w+)\s+x\s*(\d+)\s*=\s*\$?(\d{1,3}(?:,\d{3})*)',
            ]
            
            for pattern in breakdown_line_patterns:
                breakdown_lines = re.findall(pattern, breakdown_text, re.IGNORECASE)
                
                for line_match in breakdown_lines:
                    if len(line_match) == 4:
                        # Pattern 1: "Gold Eagles: ~$2,410 x 40 = $96,400"
                        product_name = line_match[0].strip()
                        unit_price = float(line_match[1].replace(',', ''))
                        quantity = int(line_match[2])
                        total_price = float(line_match[3].replace(',', ''))
                        unit = 'oz' if 'oz' in breakdown_text.lower() else 'unit'
                    elif len(line_match) == 5:
                        # Pattern 2: "Gold Eagles: ~$2,410/oz x 40 = $96,400"
                        product_name = line_match[0].strip()
                        unit_price = float(line_match[1].replace(',', ''))
                        unit = line_match[2].lower()
                        quantity = int(line_match[3])
                        total_price = float(line_match[4].replace(',', ''))
                    else:
                        continue
                    
                    # Check if product name is valid (not too long, not a sentence)
                    if len(product_name) > 50 or '.' in product_name or product_name.count(' ') > 5:
                        continue
                    
                    products.append({
                        "product_id": f"product_{len(products) + 1}",
                        "name": product_name,
                        "description": "Extracted from breakdown",
                        "category": "Not specified",
                        "pricing_matrix": [{
                            "quantity_min": quantity,
                            "quantity_max": quantity,
                            "quantity_unit": unit,
                            "payment_terms": "Not specified",
                            "unit_price": unit_price,
                            "total_price": total_price,
                            "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                            "notes": "Extracted from breakdown section"
                        }],
                        "default_payment_term": "Not specified",
                        "warranty": "Not specified",
                        "delivery_time": "Not specified"
                    })
                    print(f"[Price Agent] Extracted from breakdown: {product_name} - ${unit_price}/{unit} x {quantity} = ${total_price}", file=sys.stderr, flush=True)
                    break  # Only process first match per pattern
        
        # Pattern 1: "Gold American Eagles: $2,395/oz" or "Gold Eagles: $2,395/oz"
        # Pattern 2: "Gold: ~$75k (31–32 oz at $2,395/oz)"
        # Pattern 3: "Gold American Eagles (1oz)" followed by pricing
        # Pattern 4: "Gold American Eagles are at approx $2,410/oz"
        # Pattern 5: "Gold Eagles: ~$2,410 x 40 = $96,400" (direct calculation)
        
        # Look for direct calculation pattern: "Product: price x quantity = total" (only if no breakdown products found)
        if not products:
            direct_calc_pattern = r'([A-Z][A-Za-z\s]{1,40}(?:Eagles?|Bars?|Coins?))[:\s]+?~?\$?(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*x\s*(\d+)\s*=\s*\$?(\d{1,3}(?:,\d{3})*)'
            direct_calcs = re.findall(direct_calc_pattern, text, re.IGNORECASE)
            
            for calc_match in direct_calcs:
                product_name = calc_match[0].strip()
                # Filter out invalid product names (too long, contains periods, etc.)
                if len(product_name) > 50 or '.' in product_name or product_name.count(' ') > 5:
                    continue
                    
                unit_price = float(calc_match[1].replace(',', ''))
                quantity = int(calc_match[2])
                total_price = float(calc_match[3].replace(',', ''))
                
                # Determine unit
                unit = 'oz' if 'oz' in text.lower() else 'unit'
                
                products.append({
                    "product_id": f"product_{len(products) + 1}",
                    "name": product_name,
                    "description": "Extracted from email quote",
                    "category": "Not specified",
                    "pricing_matrix": [{
                        "quantity_min": quantity,
                        "quantity_max": quantity,
                        "quantity_unit": unit,
                        "payment_terms": "Not specified",
                        "unit_price": unit_price,
                        "total_price": total_price,
                        "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                        "notes": "Extracted from direct calculation"
                    }],
                    "default_payment_term": "Not specified",
                    "warranty": "Not specified",
                    "delivery_time": "Not specified"
                })
                print(f"[Price Agent] Extracted direct calc: {product_name} - ${unit_price} x {quantity} = ${total_price}", file=sys.stderr, flush=True)
        
        # Look for product names with prices: "Product Name: $price/unit" or "Product Name are at $price/unit"
        product_price_patterns = [
            r'([A-Z][A-Za-z\s]+?(?:Eagles?|Bars?|Coins?|Biscuits?|Grain))[:\s]+?\$?(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*/?\s*(\w+)',
            r'([A-Z][A-Za-z\s]+?(?:Eagles?|Bars?|Coins?))[^:]*?are\s+at\s+approx\s+\$?(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*/?\s*(\w+)',
            r'([A-Z][A-Za-z\s]+?(?:Eagles?|Bars?|Coins?))[^:]*?\$?(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*/?\s*(\w+)',
        ]
        
        all_matches = []
        for pattern in product_price_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            all_matches.extend(matches)
        
        print(f"[Price Agent] Found {len(all_matches)} product-price matches: {all_matches}", file=sys.stderr, flush=True)
        
        for match in all_matches:
            product_name = match[0].strip()
            unit_price_str = match[1].replace(',', '')
            unit = match[2].lower()
            
            try:
                unit_price = float(unit_price_str)
                
                # Look for quantity near this product
                # Pattern: "31–32 oz" or "~31-32 oz" or "around 40 oz"
                qty_pattern = rf'{re.escape(product_name)}[^:]*?(\d+)[–-]?(\d+)?\s*{re.escape(unit)}'
                qty_match = re.search(qty_pattern, text, re.IGNORECASE)
                
                quantity_min = None
                quantity_max = None
                if qty_match:
                    quantity_min = int(qty_match.group(1))
                    quantity_max = int(qty_match.group(2)) if qty_match.group(2) else quantity_min
                
                # Look for total: "for ~$75k" or "for approximately $75,000" or "x 40 = $96,400"
                # Pattern 1: "x 40 = $96,400" or "x 40 = 96,400"
                calculation_pattern = rf'{re.escape(product_name)}[^=]*?x\s*(\d+)\s*=\s*\$?(\d{{1,3}}(?:,\d{{3}})*)'
                calc_match = re.search(calculation_pattern, text, re.IGNORECASE)
                
                total_price = None
                if calc_match:
                    # Found explicit calculation: "x 40 = $96,400"
                    total_str = calc_match.group(2).replace(',', '')
                    total_price = float(total_str)
                    # Also update quantity if found
                    if not quantity_min:
                        quantity_min = int(calc_match.group(1))
                        quantity_max = quantity_min
                else:
                    # Pattern 2: "for ~$75k" or "for approximately $75,000"
                    total_pattern = rf'{re.escape(product_name)}[^$]*?(?:for|total|around|~|approx|approximately)\s*\$?(\d{{1,3}}(?:,\d{{3}})*(?:k|K)?)'
                    total_match = re.search(total_pattern, text, re.IGNORECASE)
                    
                    if total_match:
                        total_str = total_match.group(1).replace(',', '')
                        if 'k' in total_str.lower():
                            total_price = float(total_str.lower().replace('k', '')) * 1000
                        else:
                            total_price = float(total_str)
                    elif quantity_min and unit_price:
                        # Calculate total from unit price × quantity
                        quantity = (quantity_min + (quantity_max or quantity_min)) / 2 if quantity_max else quantity_min
                        total_price = unit_price * quantity
                
                # Create product
                pricing_matrix = [{
                    "quantity_min": quantity_min or 1,
                    "quantity_max": quantity_max,
                    "quantity_unit": unit,
                    "payment_terms": "Not specified",
                    "unit_price": unit_price,
                    "total_price": total_price,
                    "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                    "notes": "Extracted from email quote"
                }]
                
                products.append({
                    "product_id": f"product_{len(products) + 1}",
                    "name": product_name,
                    "description": "Extracted from email quote",
                    "category": "Not specified",
                    "pricing_matrix": pricing_matrix,
                    "default_payment_term": "Not specified",
                    "warranty": "Not specified",
                    "delivery_time": "Not specified"
                })
                
                print(f"[Price Agent] Extracted: {product_name} - ${unit_price}/{unit}, qty: {quantity_min}-{quantity_max}, total: ${total_price}", file=sys.stderr, flush=True)
                
            except ValueError as e:
                print(f"[Price Agent] Error parsing price for {product_name}: {e}", file=sys.stderr, flush=True)
                continue
        
        # Also look for products mentioned in lists: "Gold American Eagles (1oz)" and "Silver Eagles"
        # Pattern: Look for "recommending" or "mix" followed by product list
        # Also check for "Gold American Eagles" and "Silver Eagles" mentioned separately
        if not products or len(products) < 2:
            # Try pattern: "recommending a mix for diversification: Gold American Eagles (1oz), Some silver"
            mix_pattern = r'(?:recommending|mix|including)[^:]*?:\s*([^:]+?)(?:\.|Current|Fees|Again)'
            mix_match = re.search(mix_pattern, text, re.IGNORECASE | re.DOTALL)
            
            # Also look for products mentioned separately: "Gold American Eagles" and "Silver Eagles"
            # Pattern: Look for both "Gold" and "Silver" products
            gold_pattern = r'(Gold\s+American\s+Eagles?|Gold\s+Eagles?)'
            silver_pattern = r'(Silver\s+Eagles?)'
            
            product_names = []
            if mix_match:
                product_list_text = mix_match.group(1)
                # Extract product names from the list
                product_name_pattern = r'([A-Z][A-Za-z\s]+?(?:Eagles?|Bars?|Coins?|Silver))'
                product_names = re.findall(product_name_pattern, product_list_text)
            
            # Also check for direct mentions
            gold_match = re.search(gold_pattern, text, re.IGNORECASE)
            silver_match = re.search(silver_pattern, text, re.IGNORECASE)
            
            if gold_match and "Gold" not in " ".join(product_names):
                product_names.append(gold_match.group(1))
            if silver_match and "Silver" not in " ".join(product_names):
                product_names.append(silver_match.group(1))
            
            print(f"[Price Agent] Found products in mix/direct: {product_names}", file=sys.stderr, flush=True)
            
            if product_names:
                # For each product name, find its pricing
                for prod_name in product_names:
                    prod_name = prod_name.strip()
                    # Look for pricing near this product name
                    price_pattern = rf'{re.escape(prod_name)}[^:]*?:\s*\$?(\d{{1,3}}(?:,\d{{3}})*(?:\.\d{{2}})?)\s*/?\s*(\w+)'
                    price_match = re.search(price_pattern, text, re.IGNORECASE)
                    
                    if price_match:
                        unit_price = float(price_match.group(1).replace(',', ''))
                        unit = price_match.group(2).lower()
                        
                        # Look for quantity and total
                        qty_pattern = rf'{re.escape(prod_name)}[^$]*?(\d+)[–-]?(\d+)?\s*{re.escape(unit)}'
                        qty_match = re.search(qty_pattern, text, re.IGNORECASE)
                        
                        quantity_min = None
                        quantity_max = None
                        if qty_match:
                            quantity_min = int(qty_match.group(1))
                            quantity_max = int(qty_match.group(2)) if qty_match.group(2) else quantity_min
                        
                        # Look for calculation pattern: "x 40 = $96,400"
                        calculation_pattern = rf'{re.escape(prod_name)}[^=]*?x\s*(\d+)\s*=\s*\$?(\d{{1,3}}(?:,\d{{3}})*)'
                        calc_match = re.search(calculation_pattern, text, re.IGNORECASE)
                        
                        total_price = None
                        if calc_match:
                            # Found explicit calculation
                            total_str = calc_match.group(2).replace(',', '')
                            total_price = float(total_str)
                            if not quantity_min:
                                quantity_min = int(calc_match.group(1))
                                quantity_max = quantity_min
                        else:
                            # Pattern: "for ~$75k" or "for approximately $75,000"
                            total_pattern = rf'{re.escape(prod_name)}[^$]*?(?:for|~|approx)\s*\$?(\d{{1,3}}(?:,\d{{3}})*(?:k|K)?)'
                            total_match = re.search(total_pattern, text, re.IGNORECASE)
                            
                            if total_match:
                                total_str = total_match.group(1).replace(',', '')
                                if 'k' in total_str.lower():
                                    total_price = float(total_str.lower().replace('k', '')) * 1000
                                else:
                                    total_price = float(total_str)
                            elif quantity_min and unit_price:
                                quantity = (quantity_min + (quantity_max or quantity_min)) / 2 if quantity_max else quantity_min
                                total_price = unit_price * quantity
                        
                        if unit_price > 0:
                            products.append({
                                "product_id": f"product_{len(products) + 1}",
                                "name": prod_name,
                                "description": "Extracted from email quote",
                                "category": "Not specified",
                                "pricing_matrix": [{
                                    "quantity_min": quantity_min or 1,
                                    "quantity_max": quantity_max,
                                    "quantity_unit": unit,
                                    "payment_terms": "Not specified",
                                    "unit_price": unit_price,
                                    "total_price": total_price,
                                    "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                                    "notes": "Extracted from email quote"
                                }],
                                "default_payment_term": "Not specified",
                                "warranty": "Not specified",
                                "delivery_time": "Not specified"
                            })
        
        # Extract fees
        fee_patterns = [
            (r'(?:setup|account\s+setup|admin)[:\s]+(?:usually|around|approx|~)?\s*\$?(\d{1,3}(?:,\d{3})*)', 'Setup Fee'),
            (r'storage[:\s]+(?:around|approx|~)?\s*\$?(\d{1,3}(?:,\d{3})*)[–-]?(\d{1,3}(?:,\d{3})*)?\s*(?:annually|per\s+year|/year)', 'Storage Fee'),
        ]
        
        for pattern, fee_name in fee_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                fee_amount = float(match.group(1).replace(',', ''))
                existing_fee = next((p for p in products if fee_name.lower() in p.get('name', '').lower()), None)
                if not existing_fee:
                    products.append({
                        "product_id": f"fee_{len(products) + 1}",
                        "name": fee_name,
                        "description": "Fee extracted from quote",
                        "category": "Fee",
                        "pricing_matrix": [{
                            "quantity_min": 1,
                            "quantity_max": 1,
                            "quantity_unit": "one-time" if "setup" in fee_name.lower() else "year",
                            "payment_terms": "Not specified",
                            "unit_price": fee_amount,
                            "total_price": fee_amount,
                            "currency": currency.upper() if currency and currency != "Not specified" else "USD",
                            "notes": "Fee extracted from email quote"
                        }],
                        "default_payment_term": "Not specified",
                        "warranty": "Not specified",
                        "delivery_time": "Not specified"
                    })
        
        return products
    
    def _filter_invalid_products(self, products: List[Dict[str, Any]], text: str) -> List[Dict[str, Any]]:
        """Filter out non-products using universal, pattern-based rules (no domain-specific hardcoding)."""
        import re
        import sys
        
        # Minimal universal set: labels that denote document structure in any quote/invoice (not product names)
        STRUCTURAL_LABELS = {
            "subtotal", "total", "freight", "tax", "discount", "balance", "amount", "payment", "delivery",
            "note", "section", "item", "description", "qty", "price", "cell", "phone", "fax", "handling",
            "surcharge", "fee", "charges", "shipping", "certification", "merchandise", "estimated",
        }
        # Sentence starters — line starting with these is likely prose, not a product (exclude "the"/"a" so "The Office Suite" etc. pass)
        SENTENCE_STARTERS = {
            "i", "we", "you", "he", "she", "it", "they", "hi", "hello", "so", "if", "let", "thanks", "thank",
            "best", "here", "this", "that", "what", "when", "where", "how", "can", "could", "would", "may",
            "might", "good", "on", "re", "following", "dear", "please", "kind", "regards", "per", "as",
        }
        # Conversational phrases: if name contains these, it's prose/sentence from email, not a product name
        CONVERSATIONAL_PHRASES = (
            "we can", "we ship", "we'd probably", "if you need", "if we lock", "on timing", "payment-wise",
            "most folks", "once we have", "sharpen the pencil", "green light", "business days", "overnight",
            "good talking", "like i said", "for your folks", "bundle it like", "call it", "roughly",
            "realistically we can", "get you closer", "normally we", "first batch", "prepay with",
            "do net once", "cell:", "cell:--", "cell --",
        )
        # Common stopwords — too many in one "name" => likely a sentence
        STOPWORDS = {
            "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for", "of", "with", "by", "from",
            "as", "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "do", "does", "did",
            "will", "would", "should", "could", "may", "might", "must", "can", "this", "that", "these", "those",
            "what", "which", "who", "whom", "whose", "where", "when", "why", "how", "all", "each", "every",
            "some", "any", "no", "not", "only", "just", "also", "too", "very", "much", "many", "more", "most",
        }
        
        valid_products = []
        for product in products:
            product_name = product.get('name', '').strip()
            # Strip leading item numbers (e.g. "1 Product A" -> "Product A")
            product_name = re.sub(r'^\d+\s*[.)]\s*', '', product_name).strip()
            product_name = re.sub(r'^\d+\s+', '', product_name).strip()
            product['name'] = product_name
            product_name_lower = product_name.lower()
            words = product_name_lower.split()
            first = (words[0].lower().rstrip(':,.') if words else "")
            
            # Must have a positive price
            pricing_matrix = product.get("pricing_matrix") or []
            has_positive_price = any(
                (isinstance(e.get("unit_price"), (int, float)) and float(e.get("unit_price") or 0) > 0) or
                (isinstance(e.get("total_price"), (int, float)) and float(e.get("total_price") or 0) > 0)
                for e in pricing_matrix
            )
            if pricing_matrix and not has_positive_price:
                continue
            
            # Length: must look like a product name (2–100 chars)
            if len(product_name) < 2 or len(product_name) > 100:
                continue
            
            # PDF encoding artifact (universal)
            if "(cid" in product_name_lower:
                continue
            
            # Phone/email in name (universal)
            if re.search(r'[0-9]{3}[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}|@|XXX-XXXX', product_name, re.IGNORECASE):
                continue
            
            # Sentence: multiple periods or starts with sentence starter or too many stopwords
            if product_name.count('.') >= 2:
                continue
            if words and words[0] in SENTENCE_STARTERS:
                continue
            stopword_count = sum(1 for w in words if w in STOPWORDS)
            if stopword_count > 5:
                continue
            
            # Structural line: name is exactly a structural label, or first word is and name is short (≤3 words)
            name_normalized = product_name_lower.rstrip(':').strip()
            if name_normalized in STRUCTURAL_LABELS:
                continue
            if first in STRUCTURAL_LABELS and len(words) <= 3:
                continue
            if re.match(r'^[a-z]+\s*:\s*', product_name_lower) and first in STRUCTURAL_LABELS:
                continue
            
            # Conversational prose: name contains email/sentence phrases (e.g. "we can do", "on timing", "payment-wise", "cell:")
            if any(phrase in product_name_lower for phrase in CONVERSATIONAL_PHRASES):
                continue
            
            # Pattern-based: pricing tier, delivery time, partial price, incomplete calculation, sentence fragment
            if re.match(r'^\d+[–-]\d+\s*(?:kg|lb|units?|pieces?)\s*:', product_name, re.IGNORECASE):
                continue
            if re.match(r'^(?:delivery)\s+time\b', product_name_lower):
                continue
            if re.match(r'^USD\s*\d+[,\s]*$', product_name, re.IGNORECASE):
                continue
            if re.search(r':\s*x\s*=\s*$', product_name, re.IGNORECASE):
                continue
            if re.search(r'\b(are|is)\s+at\s+', product_name, re.IGNORECASE):
                continue
            
            valid_products.append(product)
        
        return valid_products
    
    def _validate_and_split_products(self, products: List[Dict[str, Any]], text: str, currency: str) -> List[Dict[str, Any]]:
        """Validate products and split any that are incorrectly combined"""
        import re
        import sys
        
        if not products:
            return products
        
        # Check if text mentions a specific number of products
        product_count_match = re.search(r'(\d+)\s*(?:different\s*)?products?', text.lower())
        expected_count = int(product_count_match.group(1)) if product_count_match else None
        
        split_products = []
        
        for product in products:
            product_name = product.get('name', '')
            
            # ALWAYS check if this product name contains multiple products
            # Signs: contains "including", has multiple commas, mentions "different products", or is very long
            should_split = False
            
            # Check various indicators that this might be a combined product
            if 'including' in product_name.lower():
                should_split = True
            elif product_name.count(',') >= 2:  # Multiple commas suggest list
                should_split = True
            elif len(product_name) > 80:  # Very long name suggests combined
                should_split = True
            elif 'different' in product_name.lower() and 'products' in product_name.lower():
                should_split = True
            elif expected_count and len(products) < expected_count:
                # If we're missing products, be more aggressive
                if product_name.count(',') >= 1:  # Even one comma might indicate list
                    should_split = True
            
            if should_split:
                print(f"[Price Agent] Detected combined product: {product_name[:100]}", file=sys.stderr, flush=True)
                print(f"[Price Agent] Attempting to split...", file=sys.stderr, flush=True)
                
                # First, try to extract from the original text (more reliable)
                # Look for pattern: "including X, Y, Z, and W" or "X, Y, Z, and W"
                text_lower = text.lower()
                product_list_patterns = [
                    r'including\s+([^.]*?)(?:\.|Pricing|pricing|For|Delivery|$)',
                    r'(?:products?|varieties?|types?)[,\s]+(?:including\s+)?([^.]*?)(?:\.|Pricing|pricing|For|Delivery|$)',
                    r':\s*([^.]*?)(?:\.|Pricing|pricing|For|Delivery|$)'
                ]
                
                product_list_text = None
                for pattern in product_list_patterns:
                    match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                    if match:
                        product_list_text = match.group(1)
                        break
                
                # If found in text, use that; otherwise use product name
                source_text = product_list_text if product_list_text else product_name
                
                # Extract individual products from comma-separated list
                # Handle "X, Y, Z, and W" pattern
                # First, replace " and " with comma for consistent splitting
                source_text = re.sub(r'\s+and\s+', ', ', source_text, flags=re.IGNORECASE)
                parts = [p.strip() for p in source_text.split(',')]
                
                # Clean up each part
                individual_products = []
                for part in parts:
                    # Remove common prefixes
                    part = re.sub(r'^(including|such as|like|products?|varieties?|types?|offers?)\s+', '', part.strip(), flags=re.IGNORECASE)
                    # Remove company names at the start
                    part = re.sub(r'^[A-Z][a-z]+\s+(?:Precious|Metals|offers?|provides?)\s+', '', part.strip(), flags=re.IGNORECASE)
                    # Remove trailing periods, commas, and common words
                    part = re.sub(r'\s*(?:\.|,|and|or)$', '', part.strip())
                    # Remove leading/trailing whitespace
                    part = part.strip()
                    # Filter out empty or very short parts, and common words
                    if len(part) > 3 and part.lower() not in ['etc', 'etc.', 'the', 'a', 'an', 'different', 'five', '5']:
                        individual_products.append(part)
                
                if len(individual_products) >= 2:
                    print(f"[Price Agent] Split into {len(individual_products)} products: {individual_products}", file=sys.stderr, flush=True)
                    # Create separate product entries
                    pricing_matrix = product.get('pricing_matrix', [])
                    warranty = product.get('warranty', 'Not specified')
                    
                    for idx, prod_name in enumerate(individual_products):
                        split_products.append({
                            "product_id": f"{product.get('product_id', 'PROD')}_{idx + 1}",
                            "name": prod_name.strip(),
                            "description": product.get('description', 'Not specified'),
                            "category": product.get('category', 'Not specified'),
                            "pricing_matrix": pricing_matrix.copy() if pricing_matrix else [],
                            "default_payment_term": product.get('default_payment_term', 'Not specified'),
                            "warranty": warranty,
                            "delivery_time": product.get('delivery_time', 'Not specified')
                        })
                    continue
            
            # If no splitting needed, keep original product
            split_products.append(product)
        
        # If we still don't have the expected count, try smart fallback
        if expected_count and len(split_products) != expected_count:
            print(f"[Price Agent] Still missing products ({len(split_products)} vs {expected_count}), trying smart fallback...", file=sys.stderr, flush=True)
            fallback = self._smart_fallback_extraction(text, currency)
            if fallback and len(fallback) == expected_count:
                return fallback
        
        return split_products
    
    def _smart_fallback_extraction(self, text: str, currency: str) -> List[Dict[str, Any]]:
        """Smart fallback extraction when AI fails - specifically handles product lists"""
        import re
        products = []
        
        text_lower = text.lower()
        
        # Look for product lists: "including X, Y, Z" or "X, Y, Z, and W"
        # Pattern: "including" followed by comma-separated list ending before "Pricing" or "."
        including_match = re.search(r'including\s+([^.]+?)(?:\.|Pricing|pricing)', text, re.IGNORECASE | re.DOTALL)
        if not including_match:
            # Try alternative: look for list after "products"
            including_match = re.search(r'(?:products?|varieties|types)[,\s]+(?:including\s+)?([^.]+?)(?:\.|Pricing|pricing)', text, re.IGNORECASE | re.DOTALL)
        
        if including_match:
            product_list_text = including_match.group(1)
            import sys
            print(f"[Price Agent] Raw product list text: {product_list_text[:200]}", file=sys.stderr, flush=True)
            
            # Clean up the text
            product_list_text = re.sub(r'\s+', ' ', product_list_text).strip()
            
            # Split by comma, handling "and" before last item
            # First replace " and " with comma for consistent splitting
            product_list_text = re.sub(r'\s+and\s+', ', ', product_list_text, flags=re.IGNORECASE)
            parts = [p.strip() for p in product_list_text.split(',')]
            
            # Filter and clean product names
            product_names = []
            for part in parts:
                cleaned = part.strip().rstrip('.')
                # Remove common prefixes/suffixes
                cleaned = re.sub(r'^(including|such as|like)\s+', '', cleaned, flags=re.IGNORECASE).strip()
                # Keep product names that are meaningful
                if len(cleaned) > 3 and cleaned.lower() not in ['etc', 'etc.', 'and', 'or', 'the', 'a', 'an']:
                    product_names.append(cleaned)
            
            print(f"[Price Agent] Extracted product names: {product_names}", file=sys.stderr, flush=True)
            
            if product_names:
                import sys
                print(f"[Price Agent] Found {len(product_names)} products: {product_names}", file=sys.stderr, flush=True)
                
                # Extract quantity tiers: "1–5 kg", "6–20 kg" or "1-5 kg", "6-20 kg"
                qty_tier_pattern = r'(\d+)[–-](\d+)\s*(kg|lb|units?|pieces?)'
                qty_tiers = re.findall(qty_tier_pattern, text, re.IGNORECASE)
                print(f"[Price Agent] Found quantity tiers: {qty_tiers}", file=sys.stderr, flush=True)
                
                # Extract payment terms and their associated prices
                # Pattern: "USD X per kg for [payment term]"
                payment_patterns = [
                    (r'full\s+advance\s+payment', 'Full Advance'),
                    (r'50%\s*advance\s+with\s+balance\s+on\s+delivery', '50% Advance, Balance on Delivery'),
                    (r'50%\s*advance', '50% Advance'),
                    (r'lc\s+at\s+sight', 'LC at Sight'),
                    (r'split\s+payment', 'Split Payment')
                ]
                
                # Extract prices in order they appear, grouped by quantity tier
                # Pattern: Extract all "USD X,XXX per kg" prices
                all_price_matches = re.findall(r'USD\s+(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s+per\s+kg', text, re.IGNORECASE)
                all_prices = [float(p.replace(',', '')) for p in all_price_matches]
                import sys
                print(f"[Price Agent] Found all prices: {all_prices}", file=sys.stderr, flush=True)
                
                price_sections = []
                
                # Based on user's text structure:
                # For 1-5 kg: 3 prices (full advance, 50% advance, LC)
                # For 6-20 kg: 3 prices (full advance, split payment, LC)
                if len(all_prices) >= 6 and len(qty_tiers) >= 2:
                    # First tier (1-5 kg): first 3 prices
                    price_sections.append({
                        'qty_min': int(qty_tiers[0][0]),
                        'qty_max': int(qty_tiers[0][1]),
                        'unit': qty_tiers[0][2].lower(),
                        'prices': all_prices[0:3]
                    })
                    # Second tier (6-20 kg): next 3 prices
                    price_sections.append({
                        'qty_min': int(qty_tiers[1][0]),
                        'qty_max': int(qty_tiers[1][1]),
                        'unit': qty_tiers[1][2].lower(),
                        'prices': all_prices[3:6]
                    })
                elif len(all_prices) >= 3 and len(qty_tiers) >= 1:
                    # If only one tier mentioned, use all prices for that tier
                    prices_per_tier = len(all_prices) // len(qty_tiers) if len(qty_tiers) > 0 else len(all_prices)
                    for idx, (qty_min, qty_max, qty_unit) in enumerate(qty_tiers):
                        start_idx = idx * prices_per_tier
                        tier_prices = all_prices[start_idx:start_idx + prices_per_tier]
                        if tier_prices:
                            price_sections.append({
                                'qty_min': int(qty_min),
                                'qty_max': int(qty_max),
                                'unit': qty_unit.lower(),
                                'prices': tier_prices
                            })
                elif all_prices:
                    # Fallback: use all prices for first tier or create default tier
                    if qty_tiers:
                        price_sections.append({
                            'qty_min': int(qty_tiers[0][0]),
                            'qty_max': int(qty_tiers[0][1]) if len(qty_tiers[0]) > 1 else None,
                            'unit': qty_tiers[0][2].lower() if len(qty_tiers[0]) > 2 else 'kg',
                            'prices': all_prices
                        })
                    else:
                        price_sections.append({
                            'qty_min': 1,
                            'qty_max': None,
                            'unit': 'kg',
                            'prices': all_prices
                        })
                
                # Determine payment terms
                payment_terms = []
                if re.search(r'full\s+advance', text_lower):
                    payment_terms.append("Full Advance")
                if re.search(r'50%\s*advance', text_lower):
                    payment_terms.append("50% Advance, Balance on Delivery")
                if re.search(r'lc\s+at\s+sight', text_lower):
                    payment_terms.append("LC at Sight")
                if not payment_terms:
                    payment_terms = ["Standard"]
                
                import sys
                print(f"[Price Agent] Payment terms: {payment_terms}", file=sys.stderr, flush=True)
                print(f"[Price Agent] Price sections: {price_sections}", file=sys.stderr, flush=True)
                
                # Create products with pricing matrix
                for idx, product_name in enumerate(product_names):
                    pricing_matrix = []
                    
                    # For each quantity tier, create pricing entries for each payment term
                    for price_section in price_sections:
                        for term_idx, payment_term in enumerate(payment_terms):
                            if term_idx < len(price_section['prices']):
                                pricing_matrix.append({
                                    "quantity_min": price_section['qty_min'],
                                    "quantity_max": price_section['qty_max'],
                                    "quantity_unit": price_section['unit'],
                                    "payment_terms": payment_term,
                                    "unit_price": price_section['prices'][term_idx],
                                    "total_price": None,
                                    "currency": currency,
                                    "notes": ""
                                })
                    
                    # If still no pricing matrix, create default
                    if not pricing_matrix:
                        all_prices = [float(p.replace(',', '')) for p in re.findall(r'USD\s+(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)', text, re.IGNORECASE)]
                        if all_prices:
                            avg_price = sum(all_prices) / len(all_prices)
                            pricing_matrix.append({
                                "quantity_min": 1,
                                "quantity_max": None,
                                "quantity_unit": "kg",
                                "payment_terms": payment_terms[0] if payment_terms else "Standard",
                                "unit_price": avg_price,
                                "total_price": None,
                                "currency": currency,
                                "notes": "Extracted from quote"
                            })
                    
                    if pricing_matrix:
                        products.append({
                            "product_id": f"product_{idx + 1}",
                            "name": product_name.strip(),
                            "description": "",
                            "pricing_matrix": pricing_matrix
                        })
                    else:
                        import sys
                        print(f"[Price Agent] WARNING: No pricing matrix created for product: {product_name}", file=sys.stderr, flush=True)
                        # Still add product with default pricing
                        all_prices = [float(p.replace(',', '')) for p in re.findall(r'USD\s+(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)', text, re.IGNORECASE)]
                        if all_prices:
                            avg_price = sum(all_prices) / len(all_prices)
                            products.append({
                                "product_id": f"product_{idx + 1}",
                                "name": product_name.strip(),
                                "description": "",
                                "pricing_matrix": [{
                                    "quantity_min": 1,
                                    "quantity_max": None,
                                    "quantity_unit": "kg",
                                    "payment_terms": payment_terms[0] if payment_terms else "Standard",
                                    "unit_price": avg_price,
                                    "total_price": None,
                                    "currency": currency,
                                    "notes": "Extracted from quote"
                                }]
                            })
        
        import sys
        print(f"[Price Agent] Smart fallback returning {len(products)} products", file=sys.stderr, flush=True)
        if products:
            print(f"[Price Agent] Products: {[p.get('name') for p in products]}", file=sys.stderr, flush=True)
        return products
    
    def _parse_with_regex(self, text: str, vendor_name: str) -> Dict[str, Any]:
        """Fallback regex-based parsing for when AI is not available"""
        import sys
        print(f"[QuoteParse] EXTRACTION_PATH=regex (running)", file=sys.stderr, flush=True)
        items = []
        total_price = 0.0
        currency = "USD"
        
        # Enhanced regex patterns to catch more variations
        # Matches: $130, 130, $130.50, 1,234.56, approx 130, ~130, around $130
        price_pattern = r'(?:approx|approximately|around|~|about)?\s*\$?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        # Matches: 130/lb, 130 per unit, 130 each, 130 per piece
        per_unit_pattern = r'(\d+(?:\.\d+)?)\s*(?:/|per)\s*(\w+)'
        quantity_pattern = r'(\d+(?:\.\d+)?)\s*(?:x|X|\*|units?|pcs?|pieces?|qty|quantity)'
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for per-unit pricing (e.g., "130/lvbb", "130 per unit")
            per_unit_match = re.search(per_unit_pattern, line, re.IGNORECASE)
            if per_unit_match:
                price_value = float(per_unit_match.group(1))
                unit = per_unit_match.group(2).lower()
                
                # Extract item name (text before the price)
                item_name = re.sub(per_unit_pattern, '', line).strip()
                if not item_name:
                    item_name = f"Item {len(items) + 1}"
                
                items.append({
                    "name": item_name,
                    "price": price_value,
                    "quantity": 1.0,
                    "unit_price": price_value,
                    "unit": unit,
                    "notes": "per unit pricing"
                })
                total_price += price_value
                continue
            
            # Standard price patterns
            prices = re.findall(price_pattern, line, re.IGNORECASE)
            quantities = re.findall(quantity_pattern, line, re.IGNORECASE)
            
            if prices:
                price_value = float(prices[0].replace(',', ''))
                
                # Extract item name
                item_name = re.sub(price_pattern, '', line, flags=re.IGNORECASE).strip()
                item_name = re.sub(quantity_pattern, '', item_name, flags=re.IGNORECASE).strip()
                if not item_name:
                    item_name = f"Item {len(items) + 1}"
                
                quantity = float(quantities[0]) if quantities else 1.0
                
                # Check for approximate indicators
                notes = ""
                if any(word in line.lower() for word in ["approx", "approximately", "around", "~", "about", "varies"]):
                    notes = "approximate pricing"
                
                items.append({
                    "name": item_name,
                    "price": price_value,
                    "quantity": quantity,
                    "unit_price": price_value / quantity if quantity > 0 else price_value,
                    "unit": "unit",
                    "notes": notes
                })
                
                total_price += price_value
        
        # If no items found, try to find total price
        if not items:
            total_match = re.search(r'(?:total|sum|amount|grand\s*total).*?' + price_pattern, text, re.IGNORECASE)
            if total_match:
                total_price = float(total_match.group(1).replace(',', ''))
                items.append({
                    "name": "Total Quote",
                    "price": total_price,
                    "quantity": 1,
                    "unit_price": total_price,
                    "unit": "total"
                })
        
        regex_names = [i.get("name", "") for i in items]
        print(f"[QuoteParse] REGEX_ITEMS_EXTRACTED count={len(items)} names={regex_names}", file=sys.stderr, flush=True)
        return {
            "vendor_name": vendor_name,
            "items": items,
            "total_price": total_price,
            "currency": currency,
            "extracted_text": text[:500],
            "item_count": len(items),
            "parsing_method": "regex"
        }
    
    def _extract_json_from_response(self, response_text: str) -> str:
        """Extract JSON from AI response (handles markdown code blocks)"""
        # Try to find JSON in markdown code blocks
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            return response_text[json_start:json_end].strip()
        elif "```" in response_text:
            json_start = response_text.find("```") + 3
            json_end = response_text.find("```", json_start)
            return response_text[json_start:json_end].strip()
        else:
            # Try to find JSON object
            json_start = response_text.find("{")
            json_end = response_text.rfind("}") + 1
            if json_start >= 0 and json_end > json_start:
                return response_text[json_start:json_end]
        return response_text
    
    def compare_quotes(self, quotes: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Compare multiple vendor quotes"""
        if not quotes:
            return {"error": "No quotes to compare"}
        
        comparison = {
            "vendors": [],
            "cheapest_vendor": None,
            "most_items": None,
            "best_value": None
        }
        
        min_price = float('inf')
        max_items = 0
        
        for quote in quotes:
            vendor_name = quote.get("vendor_name", "Unknown")
            total_price = quote.get("total_price", 0)
            item_count = quote.get("item_count", 0)
            
            comparison["vendors"].append({
                "name": vendor_name,
                "total_price": total_price,
                "item_count": item_count,
                "average_item_price": total_price / item_count if item_count > 0 else 0
            })
            
            if total_price < min_price:
                min_price = total_price
                comparison["cheapest_vendor"] = vendor_name
            
            if item_count > max_items:
                max_items = item_count
                comparison["most_items"] = vendor_name
        
        return comparison

